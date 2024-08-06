

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import *
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as condicao_esperada
from datetime import datetime
from time import sleep
from docx import Document
from docx.shared import Cm




def iniciar_driver():
    chrome_options = Options()
    
   

    arguments = ['--lang=pt-BR', '--start-maximized',
                '--incognito']
   
    for argument in arguments:
        chrome_options.add_argument(argument)

    

   
    chrome_options.add_experimental_option("prefs", {
       
        'download.directory_upgrade': True,
       
        'download.prompt_for_download': False,
        "profile.default_content_setting_values.notifications": 2,  
        
        "profile.default_content_setting_values.automatic_downloads": 1,
    })

    driver = webdriver.Chrome(options=chrome_options)
    wait = WebDriverWait(
        driver,
        10,
        poll_frequency=1,
        ignored_exceptions=[
            NoSuchElementException,
            ElementNotVisibleException,
            ElementNotSelectableException,
        ]
    )
    return driver,wait



driver,wait = iniciar_driver()
driver.get('https://br.tradingview.com/symbols/USDBRL/')
dados = wait.until(condicao_esperada.visibility_of_all_elements_located((By.XPATH,"//div[@class='container-pAUXADuj containerWithButton-pAUXADuj']")))

cotacao = driver.find_element(By.XPATH,"//div[@class='quotesRow-pAUXADuj']//div[@class='lastContainer-JWoJqCpY']/span[@class='last-JWoJqCpY js-symbol-last']")
cotacao_texto_puro = cotacao.text.replace(',','.')
link = driver.current_url
data = datetime.strftime(datetime.now(),'%d/%m/%Y')
cotacao_int = float(cotacao_texto_puro)
cotacao_int_final = round(cotacao_int,2)
sleep(3)

driver.save_screenshot('cotacao.png')
sleep(1)

documento = Document()
documento.add_heading(f'Cotação Atual do Dólar - {cotacao_int_final} ({data})',0)

paragrafo = documento.add_paragraph('O dólar está no valor de ')
paragrafo.add_run(f'{cotacao_int_final},').bold = True
paragrafo.add_run(f' na data {data} \n')
paragrafo.add_run(f'Valor cotado no site: {link} \n')
paragrafo.add_run('Print da cotação atual')
documento.add_picture('cotacao.png',width=Cm(15.00))
documento.add_paragraph('Cotação feita por - João Marcelo').italic = True
documento.save('teste.docx')
driver.close()

    
       